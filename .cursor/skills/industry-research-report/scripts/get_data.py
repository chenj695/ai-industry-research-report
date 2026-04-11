#!/usr/bin/env python3
import argparse
import datetime as dt
import csv
import subprocess
import json
import os
import re
import sys
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from urllib.parse import quote

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
import httpx
from openai import OpenAI
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas


MAX_TOPIC_LEN = 500
DEFAULT_MODEL = "gpt-4o-mini"
HTTP_TIMEOUT = 20.0
DEFAULT_TOPIC = "AI产业"
DEFAULT_WATCHLIST_FILE = "watchlist.json"
DEFAULT_TEMPLATE_STYLE = "industry_deep_dive"
DEFAULT_PRESET = "custom"
DEFAULT_NARRATIVE_STRENGTH = "high"

# SEC company_tickers.json uses primary symbols; map common display tickers to lookup keys.
SEC_TICKER_LOOKUP_ALIASES: Dict[str, str] = {
    "GOOG": "GOOGL",
}

# 10-digit CIK fallback when www.sec.gov ticker map is unavailable (403/blocked).
SEC_CIK_FALLBACK: Dict[str, str] = {
    "NVDA": "0001045810",
    "AMD": "0000002488",
    "MSFT": "0000789019",
    "GOOGL": "0001652044",
    "GOOG": "0001652044",
}

# 财务对比表：首列展示全称；ticker 仅用于 SEC 数据拼接（不在表内展示代码）。
US_FINANCIAL_COMPARISON_PEERS: List[Dict[str, str]] = [
    {"display": "英伟达（NVIDIA Corporation）", "ticker": "NVDA"},
    {"display": "超威半导体（Advanced Micro Devices）", "ticker": "AMD"},
    {"display": "微软（Microsoft Corporation）", "ticker": "MSFT"},
    {"display": "谷歌（Google，Alphabet Inc.）", "ticker": "GOOG"},
]


def sanitize_filename(name: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|]+", "_", name)
    cleaned = re.sub(r"\s+", "_", cleaned).strip("_")
    return cleaned[:120] or "industry_report"


def ensure_output_dirs(base_dir: Path) -> Path:
    out_dir = base_dir / "outputs"
    out_dir.mkdir(parents=True, exist_ok=True)
    return out_dir


def build_client() -> OpenAI:
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY is required.")
    base_url = os.getenv("OPENAI_BASE_URL")
    if base_url:
        return OpenAI(api_key=api_key, base_url=base_url)
    return OpenAI(api_key=api_key)


def llm_text(client: OpenAI, model: str, system_prompt: str, user_prompt: str, temperature: float = 0.3) -> str:
    """
    Return text from LLM with compatibility fallback:
    1) OpenAI Responses API
    2) OpenAI-compatible Chat Completions API (e.g. GitHub Models)
    """
    try:
        rsp = client.responses.create(
            model=model,
            temperature=temperature,
            input=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        )
        return (rsp.output_text or "").strip()
    except Exception:
        chat = client.chat.completions.create(
            model=model,
            temperature=temperature,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        )
        return ((chat.choices[0].message.content or "") if chat.choices else "").strip()


def _sec_edgar_user_agent() -> str:
    """
    SEC blocks some plain User-Agents on www.sec.gov (e.g. company_tickers.json returns 403),
    which breaks CIK mapping and yields all-N/A financial tables. Use a Mozilla-compatible string
    by default; override SEC_EDGAR_USER_AGENT with your org + email for SEC policy compliance.
    """
    ua = (os.getenv("SEC_EDGAR_USER_AGENT") or "").strip()
    if ua:
        return ua
    # www.sec.gov often 403s if User-Agent contains certain domains; keep a minimal contact URL.
    return "Mozilla/5.0 (compatible; industry-research-report/1.0; +https://example.com)"


def get_http_client() -> httpx.Client:
    return httpx.Client(
        timeout=HTTP_TIMEOUT,
        headers={
            "User-Agent": _sec_edgar_user_agent(),
            "Accept-Encoding": "gzip, deflate",
            "Accept": "application/json,text/plain,*/*",
        },
    )


def safe_get_json(client: httpx.Client, url: str) -> Dict:
    try:
        rsp = client.get(url)
        rsp.raise_for_status()
        return rsp.json()
    except Exception:
        return {}


def safe_get_text(client: httpx.Client, url: str) -> str:
    try:
        rsp = client.get(url)
        rsp.raise_for_status()
        return rsp.text
    except Exception:
        return ""


def normalize_mode(mode: str) -> str:
    v = (mode or "general").strip().lower()
    return v if v in {"general", "daily", "weekly"} else "general"


def normalize_template_style(style: str) -> str:
    v = (style or DEFAULT_TEMPLATE_STYLE).strip().lower()
    return v if v in {"industry_deep_dive", "company_initiation", "auto"} else DEFAULT_TEMPLATE_STYLE


def pick_template_style(style: str, query: str, topic: str) -> str:
    style = normalize_template_style(style)
    if style != "auto":
        return style
    text = f"{query} {topic}".lower()
    company_keywords = [
        "公司",
        "首次覆盖",
        "覆盖",
        "估值",
        "财报",
        "盈利预测",
        "管理层",
        "股权",
        "竞争优势",
        "护城河",
        "roe",
        "pe",
        "pb",
        "dcf",
        "initiation",
    ]
    for kw in company_keywords:
        if kw in text:
            return "company_initiation"
    return "industry_deep_dive"


def resolve_section_toggles(
    preset: str, include_pest: bool, include_five_forces: bool, include_segmentation: bool
) -> Tuple[bool, bool, bool]:
    p = (preset or DEFAULT_PRESET).strip().lower()
    if p == "quick":
        return False, False, False
    if p == "full":
        return True, True, True
    return include_pest, include_five_forces, include_segmentation


def normalize_narrative_strength(value: str) -> str:
    v = (value or DEFAULT_NARRATIVE_STRENGTH).strip().lower()
    return v if v in {"high", "medium"} else DEFAULT_NARRATIVE_STRENGTH


def load_watchlist(base_dir: Path, watchlist_path: str = "") -> Dict:
    default_payload = {
        "theme_keyword_map": {
            "compute": {"算力": 1.0, "gpu": 2.0, "服务器": 1.0, "idc": 1.5, "光模块": 1.8, "数据中心": 1.3, "芯片": 1.2},
            "model": {"模型": 1.0, "大模型": 1.5, "llm": 1.8, "多模态": 1.4, "agent": 1.2, "推理": 1.5, "训练": 1.2},
            "application": {"应用": 1.0, "aigc": 1.5, "saas": 1.3, "智能体": 1.4, "办公": 0.8, "营销": 0.8},
        },
        "themes": {
            "default": {
                "us_symbols_stooq": ["nvda.us", "amd.us", "tsm.us", "smci.us", "msft.us", "goog.us", "amzn.us", "meta.us"],
                "cn_hk_symbols_yahoo": {
                    "688256.SS": "寒武纪",
                    "300308.SZ": "中际旭创",
                    "601138.SS": "工业富联",
                    "0700.HK": "腾讯控股",
                    "9988.HK": "阿里巴巴-SW",
                },
            },
            "compute": {
                "us_symbols_stooq": ["nvda.us", "amd.us", "tsm.us", "smci.us", "avgo.us", "mrvl.us", "eqix.us", "dlr.us"],
                "cn_hk_symbols_yahoo": {
                    "688256.SS": "寒武纪",
                    "300308.SZ": "中际旭创",
                    "601138.SS": "工业富联",
                    "002281.SZ": "光迅科技",
                    "603083.SS": "剑桥科技",
                    "0700.HK": "腾讯控股",
                },
            },
            "model": {
                "us_symbols_stooq": ["msft.us", "goog.us", "amzn.us", "meta.us", "orcl.us", "crm.us", "adbe.us", "snow.us"],
                "cn_hk_symbols_yahoo": {
                    "0700.HK": "腾讯控股",
                    "9988.HK": "阿里巴巴-SW",
                    "9888.HK": "百度集团-SW",
                    "0241.HK": "阿里健康",
                    "600941.SS": "中国移动",
                },
            },
            "application": {
                "us_symbols_stooq": ["msft.us", "goog.us", "meta.us", "adbe.us", "crm.us", "shop.us", "now.us", "intu.us"],
                "cn_hk_symbols_yahoo": {
                    "0700.HK": "腾讯控股",
                    "9988.HK": "阿里巴巴-SW",
                    "3690.HK": "美团-W",
                    "9626.HK": "哔哩哔哩-W",
                    "002230.SZ": "科大讯飞",
                },
            },
        },
    }
    candidate = Path(watchlist_path) if watchlist_path else base_dir / DEFAULT_WATCHLIST_FILE
    try:
        if not candidate.exists():
            return default_payload
        data = json.loads(candidate.read_text(encoding="utf-8"))
        themes = data.get("themes", {})
        keyword_map = data.get("theme_keyword_map", {})
        if not isinstance(themes, dict) or not isinstance(keyword_map, dict):
            return default_payload
        if "default" not in themes:
            return default_payload
        return data
    except Exception:
        return default_payload


def _normalize_theme_bucket(bucket: Dict) -> Dict:
    us = bucket.get("us_symbols_stooq", []) if isinstance(bucket, dict) else []
    cn_hk = bucket.get("cn_hk_symbols_yahoo", {}) if isinstance(bucket, dict) else {}
    us_symbols = [str(x).strip().lower() for x in us if str(x).strip()]
    cn_hk_map = {str(k).strip(): str(v).strip() for k, v in cn_hk.items() if str(k).strip()}
    return {"us_symbols_stooq": us_symbols, "cn_hk_symbols_yahoo": cn_hk_map}


def _merge_buckets_weighted(
    themes: Dict[str, Dict], weighted_themes: List[Tuple[str, float]], default_bucket: Dict
) -> Dict:
    if not weighted_themes:
        return default_bucket
    us_weight: Dict[str, float] = {}
    cn_hk_weight: Dict[str, Tuple[str, float]] = {}
    for theme, weight in weighted_themes:
        bucket = _normalize_theme_bucket(themes.get(theme, {}))
        for sym in bucket.get("us_symbols_stooq", []):
            us_weight[sym] = us_weight.get(sym, 0.0) + weight
        for sym, name in bucket.get("cn_hk_symbols_yahoo", {}).items():
            prev = cn_hk_weight.get(sym)
            if prev:
                cn_hk_weight[sym] = (prev[0], prev[1] + weight)
            else:
                cn_hk_weight[sym] = (name, weight)

    us_sorted = [k for k, _ in sorted(us_weight.items(), key=lambda x: x[1], reverse=True)]
    cn_hk_sorted = {
        k: v[0]
        for k, v in sorted(cn_hk_weight.items(), key=lambda x: x[1][1], reverse=True)
    }
    return {"us_symbols_stooq": us_sorted, "cn_hk_symbols_yahoo": cn_hk_sorted}


def _normalize_keyword_weights(keywords: object) -> Dict[str, float]:
    # Backward compatible: supports list[str] (weight=1.0) and dict[str, number].
    result: Dict[str, float] = {}
    if isinstance(keywords, list):
        for kw in keywords:
            k = str(kw).strip().lower()
            if k:
                result[k] = 1.0
        return result
    if isinstance(keywords, dict):
        for k_raw, w_raw in keywords.items():
            k = str(k_raw).strip().lower()
            if not k:
                continue
            try:
                w = float(w_raw)
            except Exception:
                w = 1.0
            if w <= 0:
                continue
            result[k] = w
    return result


def pick_watchlist_by_topic(
    watchlist_config: Dict, text: str, force_theme: str = "", top_k: int = 2
) -> Tuple[str, Dict, Dict]:
    themes = watchlist_config.get("themes", {})
    theme_keyword_map = watchlist_config.get("theme_keyword_map", {})
    default_bucket = _normalize_theme_bucket(themes.get("default", {}))
    if force_theme and force_theme in themes:
        wl = _normalize_theme_bucket(themes.get(force_theme, {}))
        return force_theme, wl, {"selected_themes": [force_theme], "theme_scores": {force_theme: 1.0}}

    target = (text or "").lower()
    scores: Dict[str, float] = {}
    match_details: Dict[str, Dict[str, float]] = {}
    for theme, keywords in theme_keyword_map.items():
        if theme not in themes:
            continue
        score = 0.0
        hit_map: Dict[str, float] = {}
        keyword_weights = _normalize_keyword_weights(keywords)
        for kw, weight in keyword_weights.items():
            if kw in target:
                score += weight
                hit_map[kw] = weight
        scores[theme] = score
        if hit_map:
            match_details[theme] = hit_map

    ranked = [(theme, score) for theme, score in scores.items() if score > 0]
    ranked.sort(key=lambda x: x[1], reverse=True)
    if not ranked:
        return "default", default_bucket, {"selected_themes": ["default"], "theme_scores": {"default": 1.0}}

    picked = ranked[: max(1, top_k)]
    total = float(sum(x[1] for x in picked))
    weighted = [(theme, score / total) for theme, score in picked]
    merged = _merge_buckets_weighted(themes, weighted, default_bucket)
    selected = [x[0] for x in weighted]
    selected_theme = "+".join(selected)
    return selected_theme, merged, {
        "selected_themes": selected,
        "theme_scores": {theme: round(weight, 4) for theme, weight in weighted},
        "raw_weighted_hits": {theme: round(score, 4) for theme, score in picked},
        "matched_keywords": {theme: match_details.get(theme, {}) for theme in selected},
    }


def fetch_github_releases(client: httpx.Client, repo: str, limit: int = 2) -> List[Dict]:
    url = f"https://api.github.com/repos/{repo}/releases?per_page={limit}"
    data = safe_get_json(client, url)
    if not isinstance(data, list):
        return []
    out = []
    for item in data[:limit]:
        out.append(
            {
                "repo": repo,
                "tag": item.get("tag_name", ""),
                "name": item.get("name", ""),
                "published_at": item.get("published_at", ""),
                "url": item.get("html_url", ""),
            }
        )
    return out


def fetch_github_issue_velocity(client: httpx.Client, repo: str) -> Dict:
    open_url = f"https://api.github.com/search/issues?q=repo:{repo}+type:issue+state:open"
    closed_30d_url = (
        "https://api.github.com/search/issues?"
        f"q=repo:{repo}+type:issue+state:closed+closed:>={dt.date.today() - dt.timedelta(days=30)}"
    )
    open_data = safe_get_json(client, open_url)
    closed_data = safe_get_json(client, closed_30d_url)
    return {
        "repo": repo,
        "open_issues": int(open_data.get("total_count", 0)) if isinstance(open_data, dict) else 0,
        "closed_30d": int(closed_data.get("total_count", 0)) if isinstance(closed_data, dict) else 0,
    }


def fetch_hf_top_models(client: httpx.Client, limit: int = 8) -> List[Dict]:
    url = f"https://huggingface.co/api/models?sort=downloads&direction=-1&limit={limit}"
    data = safe_get_json(client, url)
    if not isinstance(data, list):
        return []
    out = []
    for item in data[:limit]:
        out.append(
            {
                "id": item.get("id", ""),
                "downloads": item.get("downloads", 0),
                "likes": item.get("likes", 0),
                "pipeline_tag": item.get("pipeline_tag", ""),
                "last_modified": item.get("lastModified", ""),
            }
        )
    return out


def fetch_google_news_rss(client: httpx.Client, query: str, limit: int = 6) -> List[Dict]:
    url = f"https://news.google.com/rss/search?q={quote(query)}&hl=zh-CN&gl=CN&ceid=CN:zh-Hans"
    text = safe_get_text(client, url)
    if not text:
        return []
    try:
        root = ET.fromstring(text)
    except Exception:
        return []
    items = []
    for item in root.findall(".//item")[:limit]:
        items.append(
            {
                "title": (item.findtext("title") or "").strip(),
                "pub_date": (item.findtext("pubDate") or "").strip(),
                "link": (item.findtext("link") or "").strip(),
                "source": (item.findtext("source") or "").strip(),
            }
        )
    return items


def fetch_arxiv_recent(client: httpx.Client, query: str, limit: int = 8) -> List[Dict]:
    q = quote((query or "artificial intelligence").strip())
    url = (
        "http://export.arxiv.org/api/query?"
        f"search_query=all:{q}&start=0&max_results={max(1, limit)}&sortBy=submittedDate&sortOrder=descending"
    )
    text = safe_get_text(client, url)
    if not text:
        return []
    try:
        root = ET.fromstring(text)
    except Exception:
        return []
    ns = {"a": "http://www.w3.org/2005/Atom"}
    out: List[Dict] = []
    for entry in root.findall("a:entry", ns)[:limit]:
        title = (entry.findtext("a:title", default="", namespaces=ns) or "").replace("\n", " ").strip()
        published = (entry.findtext("a:published", default="", namespaces=ns) or "").strip()
        summary = (entry.findtext("a:summary", default="", namespaces=ns) or "").replace("\n", " ").strip()
        link = ""
        for lk in entry.findall("a:link", ns):
            href = lk.attrib.get("href", "")
            if href:
                link = href
                break
        out.append(
            {
                "title": title,
                "published": published,
                "summary": summary[:280],
                "url": link,
            }
        )
    return out


def fetch_pypi_recent_downloads(client: httpx.Client, package: str) -> Dict:
    url = f"https://pypistats.org/api/packages/{quote(package)}/recent"
    data = safe_get_json(client, url)
    if not isinstance(data, dict):
        return {}
    payload = data.get("data", {}) if isinstance(data.get("data", {}), dict) else {}
    return {
        "package": package,
        "last_day": payload.get("last_day", 0),
        "last_week": payload.get("last_week", 0),
        "last_month": payload.get("last_month", 0),
    }


def fetch_pypi_release_meta(client: httpx.Client, package: str) -> Dict:
    url = f"https://pypi.org/pypi/{quote(package)}/json"
    data = safe_get_json(client, url)
    if not isinstance(data, dict):
        return {}
    info = data.get("info", {}) if isinstance(data.get("info", {}), dict) else {}
    releases = data.get("releases", {}) if isinstance(data.get("releases", {}), dict) else {}
    latest_ver = str(info.get("version", "") or "")
    latest_upload = ""
    files = releases.get(latest_ver, []) if latest_ver in releases and isinstance(releases.get(latest_ver), list) else []
    if files:
        latest_upload = str(files[-1].get("upload_time_iso_8601", "") or "")
    return {
        "package": package,
        "latest_version": latest_ver,
        "summary": str(info.get("summary", "") or ""),
        "latest_upload_time": latest_upload,
    }


def fetch_pypi_ecosystem(client: httpx.Client, packages: List[str]) -> List[Dict]:
    out: List[Dict] = []
    for pkg in packages:
        meta = fetch_pypi_release_meta(client, pkg)
        if not meta:
            continue
        recent = fetch_pypi_recent_downloads(client, pkg)
        out.append({**meta, **recent})
    return out


def fetch_sec_recent_filings(client: httpx.Client, tickers: List[str], per_ticker_limit: int = 3) -> List[Dict]:
    mapping_url = "https://www.sec.gov/files/company_tickers.json"
    mapping = safe_get_json(client, mapping_url)
    if not isinstance(mapping, dict):
        return []

    ticker_to_cik: Dict[str, str] = {}
    for _, row in mapping.items():
        if not isinstance(row, dict):
            continue
        t = str(row.get("ticker", "")).upper().strip()
        cik_num = row.get("cik_str", "")
        if not t:
            continue
        try:
            cik = f"{int(cik_num):010d}"
        except Exception:
            continue
        ticker_to_cik[t] = cik

    out: List[Dict] = []
    for tk in tickers:
        t = tk.upper().strip()
        lookup = SEC_TICKER_LOOKUP_ALIASES.get(t, t)
        cik = ticker_to_cik.get(lookup, "") or ticker_to_cik.get(t, "")
        if not cik:
            cik = SEC_CIK_FALLBACK.get(lookup, "") or SEC_CIK_FALLBACK.get(t, "")
        if not cik:
            continue
        sub = safe_get_json(client, f"https://data.sec.gov/submissions/CIK{cik}.json")
        if not isinstance(sub, dict):
            continue
        recent = sub.get("filings", {}).get("recent", {}) if isinstance(sub.get("filings", {}), dict) else {}
        forms = recent.get("form", []) if isinstance(recent.get("form", []), list) else []
        filing_dates = recent.get("filingDate", []) if isinstance(recent.get("filingDate", []), list) else []
        accessions = recent.get("accessionNumber", []) if isinstance(recent.get("accessionNumber", []), list) else []
        docs = recent.get("primaryDocument", []) if isinstance(recent.get("primaryDocument", []), list) else []
        count = min(len(forms), len(filing_dates), len(accessions), len(docs))
        picked = 0
        for i in range(count):
            form = str(forms[i] or "").strip()
            if form not in {"10-K", "10-Q", "8-K", "20-F", "6-K"}:
                continue
            acc = str(accessions[i] or "").replace("-", "")
            doc_name = str(docs[i] or "")
            filing_url = f"https://www.sec.gov/Archives/edgar/data/{int(cik)}/{acc}/{doc_name}" if acc and doc_name else ""
            out.append(
                {
                    "ticker": t,
                    "form": form,
                    "filing_date": str(filing_dates[i] or ""),
                    "url": filing_url,
                }
            )
            picked += 1
            if picked >= per_ticker_limit:
                break
    return out


def fetch_sec_company_financials(client: httpx.Client, tickers: List[str], lookback_points: int = 4) -> List[Dict]:
    mapping_url = "https://www.sec.gov/files/company_tickers.json"
    mapping = safe_get_json(client, mapping_url)
    if not isinstance(mapping, dict):
        return []

    ticker_to_cik: Dict[str, str] = {}
    for _, row in mapping.items():
        if not isinstance(row, dict):
            continue
        t = str(row.get("ticker", "")).upper().strip()
        cik_num = row.get("cik_str", "")
        if not t:
            continue
        try:
            cik = f"{int(cik_num):010d}"
        except Exception:
            continue
        ticker_to_cik[t] = cik

    metric_candidates = {
        "revenue": [
            "Revenues",
            "RevenueFromContractWithCustomerExcludingAssessedTax",
            "SalesRevenueNet",
            "RevenuesNetOfInterestExpense",
        ],
        "net_income": ["NetIncomeLoss", "ProfitLoss", "NetIncomeLossAvailableToCommonStockholdersBasic"],
        "gross_profit": ["GrossProfit"],
        "operating_income": ["OperatingIncomeLoss"],
        "eps_basic": ["EarningsPerShareBasic", "EarningsPerShareDiluted"],
        "cash": ["CashAndCashEquivalentsAtCarryingValue", "CashCashEquivalentsAndShortTermInvestments"],
        "r_and_d": ["ResearchAndDevelopmentExpense"],
        "capex": ["PaymentsToAcquirePropertyPlantAndEquipment", "PaymentsToAcquireProductiveAssets"],
    }

    def pick_metric_series(us_gaap: Dict, names: List[str]) -> List[Dict]:
        for name in names:
            node = us_gaap.get(name, {}) if isinstance(us_gaap, dict) else {}
            units = node.get("units", {}) if isinstance(node, dict) else {}
            for unit_key in ("USD", "USD/shares", "shares", "pure"):
                series = units.get(unit_key, []) if isinstance(units, dict) else []
                if isinstance(series, list) and series:
                    return series
            # fallback: first available unit
            if isinstance(units, dict) and units:
                for _, series_any in units.items():
                    if isinstance(series_any, list) and series_any:
                        return series_any
        return []

    def normalize_points(points: List[Dict], max_points: int) -> List[Dict]:
        cleaned: List[Dict] = []
        for p in points:
            if not isinstance(p, dict):
                continue
            cleaned.append(
                {
                    "fy": p.get("fy", ""),
                    "fp": p.get("fp", ""),
                    "form": p.get("form", ""),
                    "end": p.get("end", ""),
                    "filed": p.get("filed", ""),
                    "val": p.get("val", ""),
                }
            )
        cleaned.sort(key=lambda x: str(x.get("filed", "")), reverse=True)
        return cleaned[:max_points]

    out: List[Dict] = []
    for tk in tickers:
        display_ticker = tk.upper().strip()
        lookup = SEC_TICKER_LOOKUP_ALIASES.get(display_ticker, display_ticker)
        cik = ticker_to_cik.get(lookup, "") or ticker_to_cik.get(display_ticker, "")
        if not cik:
            cik = SEC_CIK_FALLBACK.get(lookup, "") or SEC_CIK_FALLBACK.get(display_ticker, "")
        if not cik:
            continue
        facts = safe_get_json(client, f"https://data.sec.gov/api/xbrl/companyfacts/CIK{cik}.json")
        if not isinstance(facts, dict) or not facts.get("facts"):
            continue
        facts_root = facts.get("facts", {}) if isinstance(facts.get("facts", {}), dict) else {}
        us_gaap = facts_root.get("us-gaap", {}) if isinstance(facts_root.get("us-gaap", {}), dict) else {}
        metric_payload: Dict[str, List[Dict]] = {}
        for metric_name, candidates in metric_candidates.items():
            metric_payload[metric_name] = normalize_points(pick_metric_series(us_gaap, candidates), lookback_points)
        out.append(
            {
                "ticker": display_ticker,
                "sec_lookup_ticker": lookup,
                "cik": cik,
                "entity_name": facts.get("entityName", ""),
                "currency_hint": "USD",
                "metrics": metric_payload,
            }
        )
    return out


def fetch_azure_gpu_price_indicators(client: httpx.Client) -> List[Dict]:
    base = "https://prices.azure.com/api/retail/prices"
    keywords = ["H100", "A100", "L40", "V100"]
    out: List[Dict] = []
    for kw in keywords:
        filter_expr = (
            "serviceName eq 'Virtual Machines' and "
            "priceType eq 'Consumption' and "
            f"contains(meterName,'{kw}') and "
            "armRegionName eq 'eastus'"
        )
        url = f"{base}?$filter={quote(filter_expr)}"
        data = safe_get_json(client, url)
        items = data.get("Items", []) if isinstance(data, dict) else []
        if not isinstance(items, list):
            continue
        valid = [x for x in items if isinstance(x, dict) and float(x.get("retailPrice", 0) or 0) > 0]
        if not valid:
            continue
        best = min(valid, key=lambda x: float(x.get("retailPrice", 0) or 0))
        out.append(
            {
                "gpu_keyword": kw,
                "arm_sku_name": best.get("armSkuName", ""),
                "meter_name": best.get("meterName", ""),
                "retail_price": best.get("retailPrice", 0),
                "currency": best.get("currencyCode", ""),
                "unit_of_measure": best.get("unitOfMeasure", ""),
                "region": best.get("armRegionName", ""),
            }
        )
    return out


def fetch_stooq_quotes(client: httpx.Client, symbols: List[str]) -> List[Dict]:
    if not symbols:
        return []
    q = ",".join(symbols)
    url = f"https://stooq.com/q/l/?s={q}&f=sd2t2c8ohlv&e=csv"
    text = safe_get_text(client, url)
    if not text:
        return []
    lines = [ln for ln in text.splitlines() if ln.strip()]
    if len(lines) <= 1:
        return []
    reader = csv.DictReader(lines)
    out = []
    for row in reader:
        try:
            change_pct = float(row.get("% Change", "0") or 0)
        except ValueError:
            change_pct = 0.0
        out.append(
            {
                "symbol": row.get("Symbol", ""),
                "date": row.get("Date", ""),
                "time": row.get("Time", ""),
                "close": row.get("Close", ""),
                "open": row.get("Open", ""),
                "high": row.get("High", ""),
                "low": row.get("Low", ""),
                "volume": row.get("Volume", ""),
                "change_pct": change_pct,
            }
        )
    return out


def fetch_yahoo_quotes(client: httpx.Client, symbol_name_map: Dict[str, str]) -> List[Dict]:
    symbols = list(symbol_name_map.keys())
    if not symbols:
        return []
    url = "https://query1.finance.yahoo.com/v7/finance/quote?symbols=" + quote(",".join(symbols))
    data = safe_get_json(client, url)
    result = data.get("quoteResponse", {}).get("result", []) if isinstance(data, dict) else []
    out = []
    for item in result:
        symbol = item.get("symbol", "")
        out.append(
            {
                "symbol": symbol,
                "name": symbol_name_map.get(symbol, symbol),
                "currency": item.get("currency", ""),
                "exchange": item.get("fullExchangeName", ""),
                "price": item.get("regularMarketPrice", ""),
                "change": item.get("regularMarketChange", ""),
                "change_pct": item.get("regularMarketChangePercent", 0),
                "market_time": item.get("regularMarketTime", 0),
            }
        )
    return out


def fetch_real_data(topic: str, watchlist: Dict, mode: str = "general") -> Dict:
    with get_http_client() as client:
        compute_repos = [
            "pytorch/pytorch",
            "triton-lang/triton",
            "vllm-project/vllm",
            "NVIDIA/TensorRT",
            "openxla/xla",
        ]
        model_repos = [
            "huggingface/transformers",
            "ollama/ollama",
            "ggml-org/llama.cpp",
            "openai/openai-python",
        ]

        compute_updates: List[Dict] = []
        compute_velocity: List[Dict] = []
        for repo in compute_repos:
            compute_updates.extend(fetch_github_releases(client, repo, limit=1))
            compute_velocity.append(fetch_github_issue_velocity(client, repo))

        model_updates: List[Dict] = []
        for repo in model_repos:
            model_updates.extend(fetch_github_releases(client, repo, limit=1))

        market_symbols = watchlist.get("us_symbols_stooq", [])
        market_quotes = fetch_stooq_quotes(client, market_symbols)
        market_sorted = sorted(
            market_quotes,
            key=lambda x: abs(float(x.get("change_pct", 0.0))),
            reverse=True,
        )

        cn_hk_symbol_map = watchlist.get("cn_hk_symbols_yahoo", {})
        cn_hk_quotes = fetch_yahoo_quotes(client, cn_hk_symbol_map)
        cn_hk_sorted = sorted(
            cn_hk_quotes,
            key=lambda x: abs(float(x.get("change_pct", 0) or 0)),
            reverse=True,
        )

        sec_focus_tickers = [
            "NVDA",
            "AMD",
            "MSFT",
            "GOOG",
            "AMZN",
            "META",
            "AVGO",
            "ORCL",
            "ADBE",
            "CRM",
            "TSM",
        ]
        pypi_focus_packages = [
            "torch",
            "transformers",
            "vllm",
            "openai",
            "langchain",
            "sentence-transformers",
        ]
        us_watchlist_tickers = [str(s).split(".")[0].upper() for s in market_symbols if "." in str(s)]
        sec_financial_focus = list(dict.fromkeys(us_watchlist_tickers + sec_focus_tickers))

        return {
            "generated_at": dt.datetime.now(dt.timezone.utc).isoformat(),
            "topic": topic,
            "mode": mode,
            "watchlist": watchlist,
            "compute_layer": {
                "release_updates": compute_updates,
                "issue_velocity_30d": compute_velocity,
            },
            "model_layer": {
                "release_updates": model_updates,
                "huggingface_top_models": fetch_hf_top_models(client, limit=8),
            },
            "market_layer": {
                "us_quotes": market_quotes,
                "cn_hk_quotes": cn_hk_quotes,
                "largest_daily_moves": market_sorted[:5],
                "cn_hk_largest_daily_moves": cn_hk_sorted[:5],
            },
            "news_layer": {
                "ai_general_news": fetch_google_news_rss(client, "人工智能 算力 模型", limit=6),
                "topic_news": fetch_google_news_rss(client, topic, limit=6),
            },
            "research_layer": {
                "arxiv_recent_papers": fetch_arxiv_recent(client, f"{topic} AI", limit=8),
            },
            "developer_ecosystem_layer": {
                "pypi_packages": fetch_pypi_ecosystem(client, pypi_focus_packages),
            },
            "fundamental_layer": {
                "sec_recent_filings": fetch_sec_recent_filings(client, sec_focus_tickers, per_ticker_limit=2),
                "sec_company_financials": fetch_sec_company_financials(client, sec_financial_focus, lookback_points=4),
            },
            "capital_layer": {
                "ai_funding_mna_news": fetch_google_news_rss(client, "AI 融资 并购", limit=8),
            },
            "cloud_compute_layer": {
                "azure_gpu_price_indicators": fetch_azure_gpu_price_indicators(client),
            },
        }


def extract_topic_with_llm(client: OpenAI, raw_query: str, model: str) -> str:
    prompt = (
        "你是行业关键词提取助手。"
        "请从用户输入中提取最核心的行业/产业主题，"
        "仅返回主题词，不要解释，不要加引号。"
        "如果有多个主题，请保留最核心或并列主题短语。"
    )
    topic = llm_text(client, model, prompt, raw_query, temperature=0.1)
    if not topic:
        topic = raw_query.strip()
    return topic[:MAX_TOPIC_LEN]


def slim_data_snapshot_for_llm(snap: Dict) -> Dict:
    """
    Full data_snapshot JSON can exceed small-model input limits (e.g. GitHub Models gpt-4.1 ~8k tokens
    total request). The complete snapshot is still written to *_data.json; this slim copy is for LLM prompts only.
    """

    def lim_list(obj: object, n: int) -> List:
        if not isinstance(obj, list):
            return []
        return obj[:n]

    out: Dict = {
        "generated_at": snap.get("generated_at"),
        "topic": snap.get("topic"),
        "mode": snap.get("mode"),
        "theme_mix": snap.get("theme_mix"),
        "_note": "裁剪版快照供模型引用；完整数据见输出文件 *_data.json 中的 data_snapshot。",
    }
    wl = snap.get("watchlist") if isinstance(snap.get("watchlist"), dict) else {}
    us_syms = wl.get("us_symbols_stooq") or []
    cn_map = wl.get("cn_hk_symbols_yahoo") or {}
    if isinstance(cn_map, dict):
        cn_sample = dict(list(cn_map.items())[:8])
    else:
        cn_sample = {}
    out["watchlist_summary"] = {
        "us_symbols_count": len(us_syms) if isinstance(us_syms, list) else 0,
        "us_symbols_sample": us_syms[:8] if isinstance(us_syms, list) else [],
        "cn_hk_symbols_count": len(cn_map) if isinstance(cn_map, dict) else 0,
        "cn_hk_symbols_sample": cn_sample,
    }

    cl = snap.get("compute_layer") if isinstance(snap.get("compute_layer"), dict) else {}
    out["compute_layer"] = {
        "release_updates": lim_list(cl.get("release_updates"), 6),
        "issue_velocity_30d": lim_list(cl.get("issue_velocity_30d"), 6),
    }
    ml = snap.get("model_layer") if isinstance(snap.get("model_layer"), dict) else {}
    out["model_layer"] = {
        "release_updates": lim_list(ml.get("release_updates"), 6),
        "huggingface_top_models": lim_list(ml.get("huggingface_top_models"), 8),
    }
    mk = snap.get("market_layer") if isinstance(snap.get("market_layer"), dict) else {}
    out["market_layer"] = {
        "largest_daily_moves": lim_list(mk.get("largest_daily_moves"), 6),
        "cn_hk_largest_daily_moves": lim_list(mk.get("cn_hk_largest_daily_moves"), 6),
    }
    nl = snap.get("news_layer") if isinstance(snap.get("news_layer"), dict) else {}
    out["news_layer"] = {
        "ai_general_news": lim_list(nl.get("ai_general_news"), 5),
        "topic_news": lim_list(nl.get("topic_news"), 5),
    }
    rl = snap.get("research_layer") if isinstance(snap.get("research_layer"), dict) else {}
    out["research_layer"] = {"arxiv_recent_papers": lim_list(rl.get("arxiv_recent_papers"), 5)}
    del_ec = snap.get("developer_ecosystem_layer") if isinstance(snap.get("developer_ecosystem_layer"), dict) else {}
    out["developer_ecosystem_layer"] = {"pypi_packages": lim_list(del_ec.get("pypi_packages"), 8)}
    fl = snap.get("fundamental_layer") if isinstance(snap.get("fundamental_layer"), dict) else {}
    sec_fin = fl.get("sec_company_financials") or []
    slim_fin: List[Dict] = []
    if isinstance(sec_fin, list):
        for item in sec_fin[:16]:
            if not isinstance(item, dict):
                continue
            metrics = item.get("metrics") or {}
            slim_m: Dict[str, List] = {}
            if isinstance(metrics, dict):
                for k, series in metrics.items():
                    if isinstance(series, list) and series:
                        slim_m[str(k)] = series[:1]
            slim_fin.append(
                {
                    "ticker": item.get("ticker"),
                    "entity_name": item.get("entity_name"),
                    "metrics_latest": slim_m,
                }
            )
    out["fundamental_layer"] = {
        "sec_recent_filings": lim_list(fl.get("sec_recent_filings"), 8),
        "sec_company_financials": slim_fin,
    }
    cap = snap.get("capital_layer") if isinstance(snap.get("capital_layer"), dict) else {}
    out["capital_layer"] = {"ai_funding_mna_news": lim_list(cap.get("ai_funding_mna_news"), 6)}
    ccl = snap.get("cloud_compute_layer") if isinstance(snap.get("cloud_compute_layer"), dict) else {}
    out["cloud_compute_layer"] = ccl
    return out


def generate_report_markdown(
    client: OpenAI,
    topic: str,
    model: str,
    data_snapshot: Dict,
    mode: str = "general",
    template_style: str = DEFAULT_TEMPLATE_STYLE,
    include_pest: bool = True,
    include_five_forces: bool = True,
    include_segmentation: bool = True,
    narrative_strength: str = DEFAULT_NARRATIVE_STRENGTH,
) -> Tuple[str, str]:
    mode = normalize_mode(mode)
    template_style = normalize_template_style(template_style)
    today = dt.date.today().isoformat()
    pest_section = "## 宏观环境分析（PEST）\n" if include_pest else ""
    five_forces_section = "## 行业竞争态势（波特五力）\n" if include_five_forces else ""
    segmentation_section = "## 细分市场分析（可选）\n" if include_segmentation else ""
    toggle_hint = (
        f"可选模块开关：PEST={'on' if include_pest else 'off'}，"
        f"五力={'on' if include_five_forces else 'off'}，"
        f"细分市场={'on' if include_segmentation else 'off'}。"
    )
    narrative_strength = normalize_narrative_strength(narrative_strength)
    sec_financials = (
        data_snapshot.get("fundamental_layer", {}).get("sec_company_financials", [])
        if isinstance(data_snapshot.get("fundamental_layer", {}), dict)
        else []
    )
    us_peer_displays = [p["display"] for p in US_FINANCIAL_COMPARISON_PEERS]
    available_us_with_data = []
    if isinstance(sec_financials, list):
        for item in sec_financials:
            if not isinstance(item, dict):
                continue
            tk = str(item.get("ticker", "")).upper().strip()
            if not tk:
                continue
            for peer in US_FINANCIAL_COMPARISON_PEERS:
                if peer["ticker"].upper() == tk and any(
                    item.get("metrics", {}).get(k) for k in ("revenue", "net_income")
                ):
                    available_us_with_data.append(peer["display"])
                    break
    fiscal_table_hint = (
        "重点公司财务对比要求（必须执行）：\n"
        "- 在报告中新增“重点公司财务对比表”小节，使用 Markdown 表格输出。\n"
        f"- 美国可比公司必须使用全称（禁止仅用股票代码），固定为：{'、'.join(us_peer_displays)}。\n"
        "- 表头固定为：| 公司 | 最新期别 | 营收（美元） | 净利润（美元） | 毛利（美元） | 经营利润（美元） | "
        "EPS（基本） | 现金（美元） | 研发费用（美元） | 资本开支（美元） | 数据来源与口径 |\n"
        "- 若某项缺失，填 N/A，不得编造；正文分析中提及上述公司时亦用全称，勿单独写 GOOG/NVDA 等代码。\n"
        f"- 当前快照中 SEC 侧已有数值的主体（若有）：{', '.join(available_us_with_data) or '以脚本后处理表为准'}。\n"
        "- 可比公司财务表仅限上市主体（脚本后处理表已按此口径）；非上市创新企业（如深度求索/DeepSeek、智谱、月之暗面/Kimi、MiniMax 等）"
        "不得写入该财务对比表，亦勿在正文捏造其营收利润等数字；可在「竞争格局」「技术/开源生态」「应用与商业模式」等章节做定性讨论，"
        "须标注信息来源与不确定性。\n"
    )
    narrative_hint = (
        "写作强度：medium（短平快版，结论优先、段落更短、论证简洁）"
        if narrative_strength == "medium"
        else "写作强度：high（深论证版，证据链更完整、因果展开更充分）"
    )
    narrative_rules = (
        "表达与论证要求（medium，短平快）：\n"
        "A. 每个一级章节至少1段完整论述；每段建议40-90字，先结论后要点证据。\n"
        "B. 核心判断采用“事实 -> 结论”两步表达，因果链可简写但不得缺失关键逻辑。\n"
        "C. 执行摘要输出“结论、证据、建议”三联结构。\n"
        "D. 结论与建议部分至少给出2条可执行动作，并标注优先级（高/中/低）。\n"
        "E. 周报模式下，优先写“本周变化点、原因、下周观察指标”，减少背景铺垫。\n"
        "F. 全文至少引用6个来自实时快照的具体数据点（数字/日期/涨跌幅/频次）。\n"
        "G. 每个一级章节至少包含1条带数字的事实句，避免纯观点描述。\n"
    )
    if narrative_strength == "high":
        narrative_rules = (
            "表达与论证增强要求（high，深论证，必须严格执行）：\n"
            "A. 每个「##」一级章节至少3个自然段；每段不少于4句完整中文句，单段建议120-220字；"
            "严禁用一句口号或单句列表项代替整节论证。\n"
            "B. 每个「###」小节至少2句以上展开；若用小标题下列点，每条要点须含「结论+依据+含义」mini论证，禁止仅列名词短语。\n"
            "C. 每个核心判断都采用“事实 -> 机制 -> 影响”三步表达，且每步至少一句独立陈述。\n"
            "D. 对关键结论补充反方视角或约束条件，再给最终判断，避免单边叙事。\n"
            "E. 执行摘要必须输出“结论、证据、含义、建议”四联结构；每一联至少2句，不得各用一句话敷衍。\n"
            "F. 结论与建议部分至少给出3条可执行动作，并标注优先级（高/中/低）；每条建议下再写1-2句落地做法与依据。\n"
            "G. 周报模式下，突出“本周变化点、原因、下周观察指标”；避免泛化行业科普。\n"
            "H. 公司覆盖模式下，突出“财务质量、估值框架、催化与风险对称”。\n"
            "I. 数据密度要求：全文至少引用12个来自实时数据快照的具体数据点（数值/日期/仓库名/版本号/涨跌幅/issue关闭数/下载量等）；"
            "其中至少6个必须显式出现在「数据快照」或「产业链/竞争」相关章节正文，而非仅在执行摘要一笔带过。\n"
            "J. 每个一级章节至少包含2条“数据事实句”（带明确数字或可追溯名称），并各跟1句解释其对结论的方向性影响。\n"
            "K. 预测与情景推演：乐观/基准/悲观各情景至少一个小段落（3句以上），写清关键变量、区间/量级、触发条件与失效信号。\n"
            "L. 竞争格局章节必须给出至少1个可量化集中度指标（如CR3/CR5或可替代口径），用2句以上解释计算口径与行业含义。\n"
            "M. 建议章节每条建议都要绑定对应证据（快照或公认来源），禁止“加大投入/加强合作”类无锚点空话。\n"
            "N. 禁止用未在快照或权威来源中出现的「精确市场规模整数」（如固定「5000亿美元」）作为全文核心锚点；"
            "若快照无第三方市场规模，须写区间+假设+不确定性，并明确与快照字段的对应关系。\n"
            "O. PEST 若开启：政/经/社/技每一维至少2段或一组充分展开的要点，禁止每维仅一句话。\n"
        )
    if template_style == "company_initiation":
        system_prompt = (
            "你是一名资深卖方分析师。请使用中文输出公司首次覆盖风格研究报告。"
            "输出必须是 markdown，包含：\n"
            "# 标题\n"
            "## 报告引言（目的、研究范围、核心结论摘要）\n"
            "## 执行摘要（覆盖观点、评级倾向、核心催化）\n"
            "## 数据快照（算力层->模型层->市场层）\n"
            f"{pest_section}"
            "## 公司概况与治理结构（前身、产品矩阵、历史沿革、股权结构、核心团队、激励机制）\n"
            "## 商业模式与盈利模式拆解\n"
            "## 行业空间与量价框架（量=需求，价=成本/ASP）\n"
            "## 行业发展现状（规模、增长驱动、市场结构、区域分布）\n"
            f"{segmentation_section}"
            "## 竞争格局与可比公司（同行/历史/国际）\n"
            f"{five_forces_section}"
            "## 核心竞争优势与护城河（技术、客户、成本、研发）\n"
            "## 财务质量分析（成长、盈利、现金流、周转）\n"
            "## 未来3年预测（收入、毛利、费用率、利润）\n"
            "## 估值框架与敏感性分析\n"
            "## 标杆企业分析（2-3家，可选）\n"
            "## 风险与不确定性\n"
            "## 数据来源与口径说明\n"
            "## 结论与建议\n"
            "要求：逻辑清晰，强调公司基本面，先事实后判断。"
            "对建议部分需分别给出：企业建议、投资者建议。"
            "禁止只给一句话结论，必须给出论证过程与证据链。"
            "篇幅与展开：深度报告建议总篇幅约8000-15000汉字；每个「##」章节下至少3段、每段多句展开，避免标题下仅一行结论。"
        )
    else:
        system_prompt = (
            "你是一名资深行业分析师。请使用中文输出专业、结构化行业研究报告。"
            "输出必须是 markdown，包含：\n"
            "# 标题\n"
            "## 报告引言（目的、行业定义与范围、核心结论摘要）\n"
            "## 执行摘要\n"
            "## 数据快照（算力层->模型层->市场层）\n"
            f"{pest_section}"
            "## 公司概况与治理结构（前身、产品矩阵、历史沿革、股权结构、核心团队、战略变更、财务概览）\n"
            "## 行业定义与边界\n"
            "## 行业分析：量价框架（量=下游需求/渗透率，价=成本/ASP/溢价）\n"
            "## 市场规模与增速（给出区间/假设）\n"
            "## 产业链与关键环节（上中下游）\n"
            "## 行业发展现状（增长驱动、市场结构、区域分布）\n"
            f"{segmentation_section}"
            "## 竞争格局（5力/CRn/商业模式）\n"
            f"{five_forces_section}"
            "## 竞争优势与护城河验证（产品、技术、研发、客户、成本）\n"
            "## 核心驱动因素与边际变化\n"
            "## 三维比较（同行比较、历史比较、国际比较）\n"
            "## 风险与不确定性\n"
            "## 未来3年情景推演（乐观/基准/悲观）与敏感性分析\n"
            "## 市场规模预测（3-5年，含依据）\n"
            "## 盈利预测与关键假设（收入、毛利、费用率、现金流）\n"
            "## 标杆企业分析（2-3家，可选）\n"
            "## 数据来源与口径说明\n"
            "## 结论与建议\n"
            "要求：逻辑清晰、可用于面试与业务沟通，避免空话。每一节先陈述事实数据，再给判断。"
            "禁止只给一句话结论，必须给出论证过程与证据链。"
            "篇幅与展开：深度报告建议总篇幅约8000-15000汉字；每个「##」章节至少3个自然段、每段4句以上；"
            "「宏观环境」「产业链」「竞争格局」「情景推演」等核心章不得用单句敷衍。"
        )
    mode_hint = "常规深度报告"
    if mode == "daily":
        mode_hint = "日报：重点写过去24小时更新、市场异动和次日观察点。"
    elif mode == "weekly":
        mode_hint = "周报：重点写过去7天变化、周度复盘和下周前瞻。"
    slim_snap = slim_data_snapshot_for_llm(data_snapshot)
    data_json = json.dumps(slim_snap, ensure_ascii=False)
    max_json_chars_raw = os.getenv("REPORT_DATA_SNAPSHOT_MAX_CHARS", "").strip()
    if max_json_chars_raw.isdigit():
        max_json_chars = int(max_json_chars_raw)
    else:
        ml = (model or "").lower()
        # GitHub Models 等对 gpt-4.1 类模型有较小请求体上限，默认更严。
        max_json_chars = 4500 if ("gpt-4.1" in ml or "4.1" in ml) else 14000
    if len(data_json) > max_json_chars:
        data_json = (
            data_json[: max(0, max_json_chars - 120)]
            + "\n…[数据快照 JSON 已按长度截断；完整内容见输出目录 *_data.json 的 data_snapshot 字段。]"
        )
    user_prompt = (
        f"请围绕“{topic}”生成完整行业研究报告。\n"
        f"报告模式：{mode_hint}\n"
        f"模板风格：{template_style}\n"
        f"{toggle_hint}\n"
        f"{narrative_hint}\n"
        f"报告日期：{today}\n"
        "你必须优先使用我给你的实时数据快照进行事实陈述，"
        "并在正文中明确区分“事实数据”和“基于数据的判断”。\n"
        "（下列为裁剪/可能截断后的快照；完整 JSON 已写入本机输出文件 *_data.json，勿假设未列出字段不存在。）\n\n"
        "写作规则：\n"
        "0) Markdown 小节标题请使用 ## 或 ###；避免 #### 及以下；正文与公司名勿使用股票代码简称（用全称如「谷歌（Google，Alphabet）」）。\n"
        "1) 若是行业报告，仍需给出至少1-2个代表性公司的小节分析作为落地样本。\n"
        "2) 行业分析必须围绕“量价”展开，并给出可验证先导指标。\n"
        "3) 预测部分必须明确关键假设、变量方向和主要风险触发条件。\n"
        "4) 禁止使用未经验证的自媒体二手数据作为核心论据。\n"
        "5) 必须体现章节之间的因果链：环境 -> 现状 -> 竞争 -> 趋势/预测 -> 建议。\n"
        "6) 尽量提供表格化摘要（如CRn、CAGR、三情景假设）。\n"
        "7) 反单薄：禁止在「##」标题下只写一句话或只列无解释的关键词；"
        "若某节材料不足，应明确写数据缺口并基于快照可做部分展开，而不是用泛泛行业常识凑字数。\n\n"
        "数据与论证强化规则（必须执行）：\n"
        "a) 先写“事实数据”，再写“基于数据的判断”，两者必须成对出现。\n"
        "b) 关键段落禁止只写观点，至少包含1个数字型证据（价格、涨跌幅、发布频率、issue关闭数、下载量等）。\n"
        "c) 对主要结论至少补充1条反向证据或不确定性约束，再给最终判断。\n"
        "d) 在执行摘要末尾增加“证据强度自评”（高/中/低）与原因。\n"
        "e) 对缺失数据要明确写“数据缺口与替代口径”，不得编造。\n\n"
        f"{fiscal_table_hint}\n"
        f"{narrative_rules}\n"
        f"实时数据快照如下：\n{data_json}\n"
    )
    llm_temp = 0.42 if narrative_strength == "high" else 0.35
    text = llm_text(client, model, system_prompt, user_prompt, temperature=llm_temp)
    if not text:
        text = f"# {topic}行业研究报告\n\n## 执行摘要\n暂无内容。"
    title_match = re.search(r"^#\s+(.+)$", text, flags=re.MULTILINE)
    title = title_match.group(1).strip() if title_match else f"{topic}行业研究报告"
    if mode == "daily" and "日报" not in title:
        title = f"{today} {topic}产业日报"
    if mode == "weekly" and "周报" not in title:
        iso_week = dt.date.today().isocalendar().week
        title = f"{dt.date.today().year}W{iso_week:02d} {topic}产业周报"
    return title, text


def write_markdown(md_path: Path, content: str) -> None:
    md_path.write_text(content, encoding="utf-8")


def write_json(json_path: Path, payload: Dict) -> None:
    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def write_docx(docx_path: Path, title: str, markdown_text: str) -> None:
    doc = Document()

    # Global document typography: Chinese Songti + English Times New Roman.
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal_style.font.size = Pt(12)  # 小四
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal_style.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)

    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.first_line_indent = Pt(0)
    title_para.paragraph_format.space_before = Pt(12)
    title_para.paragraph_format.space_after = Pt(12)
    for run in title_para.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(16)  # 三号
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

    def clean_inline_md(text: str) -> str:
        text = text.replace("**", "").replace("__", "").replace("`", "")
        text = text.strip()
        # 去掉行首误留的 markdown 标题井号（表格单元格、正文偶发）
        text = re.sub(r"^#{1,6}\s+", "", text)
        return text.strip()

    def is_md_table_line(text: str) -> bool:
        t = text.strip()
        return t.startswith("|") and "|" in t[1:]

    def split_md_row(text: str) -> List[str]:
        t = text.strip().strip("|")
        return [c.strip() for c in t.split("|")]

    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        striped = lines[i].strip()

        if is_md_table_line(striped):
            table_lines: List[str] = []
            while i < len(lines) and is_md_table_line(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            if table_lines:
                header = split_md_row(table_lines[0])
                body_lines = table_lines[1:]
                sep_pattern = re.compile(r"^\|\s*[-:| ]+\|?\s*$")
                body_lines = [ln for ln in body_lines if not sep_pattern.match(ln)]
                table = doc.add_table(rows=1, cols=max(1, len(header)))
                table.style = "Table Grid"
                for c_idx, val in enumerate(header):
                    p = table.rows[0].cells[c_idx].paragraphs[0]
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    run = p.add_run(clean_inline_md(val))
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(12)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                for row in body_lines:
                    vals = split_md_row(row)
                    r = table.add_row()
                    for c_idx in range(len(header)):
                        cell_val = vals[c_idx] if c_idx < len(vals) else ""
                        p = r.cells[c_idx].paragraphs[0]
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                        run = p.add_run(clean_inline_md(cell_val))
                        run.font.name = "Times New Roman"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 0, 0)
            continue

        i += 1
        if not striped:
            para = doc.add_paragraph("")
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            para.paragraph_format.first_line_indent = Pt(24)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            continue
        if striped in {"---", "***", "___"}:
            continue
        if striped.startswith("###### "):
            para = doc.add_heading(clean_inline_md(striped[7:]), level=4)
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            for run in para.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(12)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif striped.startswith("##### "):
            para = doc.add_heading(clean_inline_md(striped[6:]), level=4)
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            for run in para.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(12)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif striped.startswith("#### "):
            para = doc.add_heading(clean_inline_md(striped[5:]), level=4)
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(4)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            para.paragraph_format.keep_with_next = True
            for run in para.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(13)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif striped.startswith("### "):
            para = doc.add_heading(clean_inline_md(striped[4:]), level=3)
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            para.paragraph_format.keep_with_next = True
            for run in para.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(14)  # 四号
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif striped.startswith("## "):
            para = doc.add_heading(clean_inline_md(striped[3:]), level=2)
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after = Pt(8)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            para.paragraph_format.keep_with_next = True
            for run in para.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(15)  # 小三
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif striped.startswith("# "):
            continue
        else:
            para = doc.add_paragraph(clean_inline_md(striped))
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            para.paragraph_format.first_line_indent = Pt(24)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.widow_control = True
            for run in para.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(12)  # 小四
                run.font.color.rgb = RGBColor(0, 0, 0)
    doc.save(str(docx_path))


def write_pdf(pdf_path: Path, title: str, markdown_text: str) -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    c = canvas.Canvas(str(pdf_path), pagesize=A4)
    width, height = A4
    x = 40
    y = height - 40
    line_height = 16
    c.setFont("STSong-Light", 14)
    c.drawString(x, y, title)
    y -= 24
    c.setFont("STSong-Light", 11)

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip().replace("#", "").strip()
        if not line:
            y -= line_height
        else:
            parts = [line[i : i + 45] for i in range(0, len(line), 45)]
            for p in parts:
                if y < 40:
                    c.showPage()
                    c.setFont("STSong-Light", 11)
                    y = height - 40
                c.drawString(x, y, p)
                y -= line_height
        if y < 40:
            c.showPage()
            c.setFont("STSong-Light", 11)
            y = height - 40
    c.save()


def build_html_share(share_path: Path, title: str, markdown_text: str) -> None:
    body = markdown_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    html = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>{title}</title>
  <style>
    body {{ font-family: -apple-system, BlinkMacSystemFont, Segoe UI, Roboto, sans-serif; margin: 32px; line-height: 1.75; }}
    pre {{ white-space: pre-wrap; word-break: break-word; }}
  </style>
</head>
<body>
  <h1>{title}</h1>
  <pre>{body}</pre>
</body>
</html>
"""
    share_path.write_text(html, encoding="utf-8")


def make_truncated_text(markdown_text: str, limit: int = 600) -> str:
    plain = re.sub(r"^#{1,6}\s*", "", markdown_text, flags=re.MULTILINE).strip()
    plain = re.sub(r"\n{3,}", "\n\n", plain)
    if len(plain) <= limit:
        return plain
    return plain[:limit].rstrip() + "..."


def _fmt_num(v: object) -> str:
    try:
        if v is None or v == "":
            return "N/A"
        n = float(v)
        if abs(n) >= 1000:
            return f"{n:,.0f}"
        return f"{n:.2f}".rstrip("0").rstrip(".")
    except Exception:
        return str(v) if str(v).strip() else "N/A"


def _latest_metric(metrics: Dict, key: str) -> Tuple[str, str]:
    series = metrics.get(key, []) if isinstance(metrics, dict) else []
    if not isinstance(series, list) or not series:
        return "N/A", "N/A"
    p = series[0] if isinstance(series[0], dict) else {}
    period = f"{p.get('fy', '')}{p.get('fp', '')}".strip() or str(p.get("end", "") or "N/A")
    return _fmt_num(p.get("val", "")), period


def build_financial_comparison_table(data_snapshot: Dict) -> str:
    sec_financials = (
        data_snapshot.get("fundamental_layer", {}).get("sec_company_financials", [])
        if isinstance(data_snapshot.get("fundamental_layer", {}), dict)
        else []
    )
    by_ticker: Dict[str, Dict] = {}
    if isinstance(sec_financials, list):
        for item in sec_financials:
            if not isinstance(item, dict):
                continue
            tk = str(item.get("ticker", "")).upper().strip()
            if tk:
                by_ticker[tk] = item

    lines = [
        "## 重点公司财务对比表",
        "",
        "说明：下表为美国上市公司 SEC EDGAR 合并报表口径（美元）。",
        "",
        "| 公司 | 最新期别 | 营收（美元） | 净利润（美元） | 毛利（美元） | 经营利润（美元） | EPS（基本） | 现金（美元） | 研发费用（美元） | 资本开支（美元） | 数据来源与口径 |",
        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
    ]
    for peer in US_FINANCIAL_COMPARISON_PEERS:
        tk = peer["ticker"].upper().strip()
        display = peer["display"]
        item = by_ticker.get(tk, {})
        metrics = item.get("metrics", {}) if isinstance(item, dict) else {}
        revenue, p_revenue = _latest_metric(metrics, "revenue")
        net_income, p_net = _latest_metric(metrics, "net_income")
        gross_profit, p_gross = _latest_metric(metrics, "gross_profit")
        op_income, p_op = _latest_metric(metrics, "operating_income")
        eps, p_eps = _latest_metric(metrics, "eps_basic")
        cash, p_cash = _latest_metric(metrics, "cash")
        rnd, p_rnd = _latest_metric(metrics, "r_and_d")
        capex, p_capex = _latest_metric(metrics, "capex")
        period = next((p for p in [p_revenue, p_net, p_gross, p_op, p_eps, p_cash, p_rnd, p_capex] if p != "N/A"), "N/A")
        if not item:
            source = "SEC 数据未获取（网络、User-Agent 或 CIK 映射失败时可出现）"
        elif tk == "GOOG":
            source = "SEC EDGAR（Alphabet Inc. 合并报表，美元；谷歌为旗下业务品牌）"
        else:
            source = "SEC EDGAR（美国上市公司合并报表，美元）"
        lines.append(
            f"| {display} | {period} | {revenue} | {net_income} | {gross_profit} | {op_income} | {eps} | {cash} | {rnd} | {capex} | {source} |"
        )
    lines.append("")
    return "\n".join(lines)


def ensure_financial_table(markdown_text: str, data_snapshot: Dict) -> str:
    table_block = build_financial_comparison_table(data_snapshot)
    marker = "## 重点公司财务对比表"
    if marker in markdown_text:
        pattern = re.compile(r"## 重点公司财务对比表\s*\n(?:.*\n)*?(?=\n## [^#]|$)", re.MULTILINE | re.DOTALL)
        return pattern.sub(table_block.rstrip() + "\n\n", markdown_text, count=1)

    insert_key = "## 结论与建议"
    if insert_key in markdown_text:
        return markdown_text.replace(insert_key, f"{table_block}\n{insert_key}", 1)
    return markdown_text.rstrip() + "\n\n" + table_block + "\n"


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate industry research report.")
    parser.add_argument("--query", required=False, default=DEFAULT_TOPIC, help="Free-form user query or topic")
    parser.add_argument(
        "--mode",
        required=False,
        default="general",
        choices=["general", "daily", "weekly"],
        help="Report mode",
    )
    parser.add_argument(
        "--open-share",
        action="store_true",
        help="Open generated share html in default browser",
    )
    parser.add_argument(
        "--watchlist",
        required=False,
        default="",
        help="Optional custom watchlist JSON path",
    )
    parser.add_argument(
        "--theme",
        required=False,
        default="auto",
        help="Theme for watchlist: auto/default/compute/model/application",
    )
    parser.add_argument(
        "--mix-top-k",
        required=False,
        type=int,
        default=2,
        help="When theme=auto, mix top-k themes by keyword hit weights",
    )
    parser.add_argument(
        "--template-style",
        required=False,
        default=DEFAULT_TEMPLATE_STYLE,
        choices=["industry_deep_dive", "company_initiation", "auto"],
        help="Report template style",
    )
    parser.add_argument("--include-pest", dest="include_pest", action="store_true", default=True, help="Include PEST section")
    parser.add_argument("--no-include-pest", dest="include_pest", action="store_false", help="Disable PEST section")
    parser.add_argument("--include-five-forces", dest="include_five_forces", action="store_true", default=True, help="Include Five Forces section")
    parser.add_argument("--no-include-five-forces", dest="include_five_forces", action="store_false", help="Disable Five Forces section")
    parser.add_argument("--include-segmentation", dest="include_segmentation", action="store_true", default=True, help="Include segmentation section")
    parser.add_argument("--no-include-segmentation", dest="include_segmentation", action="store_false", help="Disable segmentation section")
    parser.add_argument(
        "--preset",
        required=False,
        default=DEFAULT_PRESET,
        choices=["custom", "quick", "full"],
        help="Section preset: quick=short version, full=all sections, custom=manual toggles",
    )
    parser.add_argument(
        "--narrative-strength",
        required=False,
        default=DEFAULT_NARRATIVE_STRENGTH,
        choices=["high", "medium"],
        help="Narrative depth: medium=short concise, high=deep argumentative",
    )
    args = parser.parse_args()

    query = (args.query or "").strip()
    mode = normalize_mode(args.mode)
    requested_template_style = normalize_template_style(args.template_style)
    narrative_strength = normalize_narrative_strength(args.narrative_strength)
    include_pest, include_five_forces, include_segmentation = resolve_section_toggles(
        args.preset, args.include_pest, args.include_five_forces, args.include_segmentation
    )
    if not query:
        query = DEFAULT_TOPIC
    if len(query) > MAX_TOPIC_LEN:
        print("ERROR_TOPIC_TOO_LONG", file=sys.stderr)
        return 2

    try:
        base_dir = Path(__file__).resolve().parents[1]
        out_dir = ensure_output_dirs(base_dir)
        watchlist_config = load_watchlist(base_dir, args.watchlist)
        client = build_client()
        model = os.getenv("OPENAI_MODEL", DEFAULT_MODEL)
        topic = extract_topic_with_llm(client, query, model)
        resolved_template_style = pick_template_style(requested_template_style, query, topic)
        force_theme = "" if (args.theme or "auto") == "auto" else (args.theme or "").strip().lower()
        selected_theme, watchlist, mix_info = pick_watchlist_by_topic(
            watchlist_config,
            text=f"{query} {topic}",
            force_theme=force_theme,
            top_k=max(1, args.mix_top_k),
        )
        data_snapshot = fetch_real_data(topic, watchlist=watchlist, mode=mode)
        data_snapshot["theme_mix"] = mix_info
        title, md_text = generate_report_markdown(
            client,
            topic,
            model,
            data_snapshot,
            mode=mode,
            template_style=resolved_template_style,
            include_pest=include_pest,
            include_five_forces=include_five_forces,
            include_segmentation=include_segmentation,
            narrative_strength=narrative_strength,
        )
        md_text = ensure_financial_table(md_text, data_snapshot)

        ts = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
        safe = sanitize_filename(title)
        stem = f"{ts}_{safe}"

        md_path = out_dir / f"{stem}.md"
        docx_path = out_dir / f"{stem}.docx"
        pdf_path = out_dir / f"{stem}.pdf"
        json_path = out_dir / f"{stem}.json"
        data_path = out_dir / f"{stem}_data.json"
        html_path = out_dir / f"{stem}.html"

        write_markdown(md_path, md_text)
        write_docx(docx_path, title, md_text)
        write_pdf(pdf_path, title, md_text)
        build_html_share(html_path, title, md_text)
        write_json(data_path, data_snapshot)

        payload = {
            "title": title,
            "topic": topic,
            "mode": mode,
            "preset": args.preset,
            "narrative_strength": narrative_strength,
            "template_style": resolved_template_style,
            "template_style_requested": requested_template_style,
            "include_pest": include_pest,
            "include_five_forces": include_five_forces,
            "include_segmentation": include_segmentation,
            "selected_theme": selected_theme,
            "theme_mix": mix_info,
            "watchlist_path": (Path(args.watchlist).as_posix() if args.watchlist else (base_dir / DEFAULT_WATCHLIST_FILE).as_posix()),
            "truncated_text": make_truncated_text(md_text),
            "pdf_output_path": pdf_path.as_posix(),
            "docx_output_path": docx_path.as_posix(),
            "md_output_path": md_path.as_posix(),
            "json_output_path": json_path.as_posix(),
            "data_output_path": data_path.as_posix(),
            "share_url": html_path.resolve().as_uri(),
            "data_snapshot": data_snapshot,
        }
        write_json(json_path, payload)
        if args.open_share:
            try:
                subprocess.run(["cmd", "/c", "start", "", str(html_path.resolve())], check=False)
            except Exception:
                pass
        print(json.dumps(payload, ensure_ascii=False))
        return 0
    except Exception as exc:
        print("报告生成服务暂时不可用，请稍后重试。", file=sys.stderr)
        if os.getenv("REPORT_DEBUG", "").strip().lower() in {"1", "true", "yes"}:
            print(f"{type(exc).__name__}: {exc}", file=sys.stderr)
            import traceback

            traceback.print_exc(file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
